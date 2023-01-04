from xml.etree.ElementTree import tostring
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import csv


def ProcessCSV(filename):
    sends = []
    with open(filename,'r') as file:
        reader = csv.reader(file)
        for row in reader:
            sends.append(row)
    return sends

def loadReplacements(config = "replacementConfig.txt"):
    replacementFile = open(config,'r')
    lines = replacementFile.readlines()

    
    replacements = {}
    
    for item in lines:
        p = item.split(",",1)
        replacements[p[0]] = p[1]
    return replacements

rp = 'FL2F-combined-consent.docx'
mda = 'MDA Agreement.docx'


import sys
import subprocess
import re

def libreoffice_exec():
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    else:
        print(sys.platform)
    return 'libreoffice'

def convert_to(folder,source,timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]
 
    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)

    filename = re.search('-> (.*?) using filter', process.stdout.decode())
    print(filename)
    return filename.group(1)

def processDocument(docFile,replacements,participants):
    doc = Document(docFile)
    today = date.today()
    for p in doc.paragraphs:
        for replacement in replacements.keys():
            if p.text.find(replacement) != -1:
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                rScript = replacements[replacement]
                try:
                    
                    globalsParameter = {'__builtins__' : None,'strftime':date.strftime}
                    localsParameter = {'p': p, 'participants': participants, 'today':today}
                    exec(rScript,globalsParameter,localsParameter)
                except: 
                    print("Malformed Replacement: ", rScript)
    
    doc.save("MDA Agreement.docx")
    newFile = convert_to('.',"MDA Agreement.docx")
    print(newFile)
   

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

def SendEmail(email,pdfs = None):
    gmailAccount = "devin@atkin.engineer"
    password = "bnrdftxpotnonaxh"
    sender_address = "theteam@fl2f.ca"

    mail_content = "Hello, \nPlease Fill and Sign the Attached Documents for the upcoming FL2F Workshop.\nThank You"

    session = smtplib.SMTP('smtp.gmail.com', 587) #use gmail with port
    session.starttls() #enable security
    session.login(gmailAccount, password) #login with mail_id and password

    message = MIMEMultipart()
    message['To'] = email
    message['From'] = sender_address
    message['Subject'] = 'MDA and Recording Permissions Form'   #The subject line
    message.attach(MIMEText(mail_content, 'plain'))
    if pdfs is not None:
        for pdf in pdfs:
            pdfFile = MIMEApplication(open(pdf,'rb').read())
            pdfFile.add_header('Content-Disposition','attachment',filename= pdf)
            message.attach(pdfFile)

    text = message.as_string()
    session.sendmail(sender_address, email, text)

    session.quit()
